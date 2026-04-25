"""Run this once to generate WeJaWi_SOP.docx in the project root."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

DOC_OUT = "WeJaWi_SOP.docx"

# ── colour palette ────────────────────────────────────────────────────────────
PURPLE     = RGBColor(0x7C, 0x3A, 0xED)   # brand primary
DEEP       = RGBColor(0x0B, 0x0B, 0x14)   # near-black
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xF5, 0xF3, 0xFF)
MID_GREY   = RGBColor(0x5B, 0x58, 0x70)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_page_break(doc):
    doc.add_page_break()


def heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = PURPLE if level <= 2 else DEEP
    return p


def body(doc, text: str):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def bullet(doc, text: str):
    doc.add_paragraph(text, style="List Bullet")


def step(doc, num: int, text: str):
    doc.add_paragraph(f"{num}. {text}", style="List Number")


# ── build document ────────────────────────────────────────────────────────────
doc = Document()

# ── global styles ─────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Helvetica Neue"
style.font.size = Pt(10.5)
style.font.color.rgb = DEEP

for i in range(1, 5):
    h = doc.styles[f"Heading {i}"]
    h.font.name = "Helvetica Neue"
    h.font.bold = True
    h.font.color.rgb = PURPLE

doc.styles["List Bullet"].font.size = Pt(10.5)
doc.styles["List Number"].font.size = Pt(10.5)

# page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)


# ════════════════════════════════════════════════════════════════════════════════
# PAGE 1 — COVER
# ════════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("WEJAWI STUDIO")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = PURPLE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Standard Operating Procedure")
run.font.size = Pt(18)
run.font.color.rgb = MID_GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\nContent Creation Automation Desktop Application")
run.font.size = Pt(13)
run.font.color.rgb = DEEP

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"\n\nVersion 1.0  ·  {datetime.date.today().strftime('%B %Y')}")
run.font.size = Pt(10)
run.font.color.rgb = MID_GREY

add_page_break(doc)


# ════════════════════════════════════════════════════════════════════════════════
# PAGE 2 — TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, "Contents", 1)
toc_items = [
    ("1", "Overview & Purpose", "3"),
    ("2", "System Requirements", "3"),
    ("3", "Feature Set", "4"),
    ("4", "Standard Operating Procedures", "6"),
    ("  4.1", "Browse (Home)", "6"),
    ("  4.2", "Transcribe", "6"),
    ("  4.3", "Captions", "7"),
    ("  4.4", "Stitch Up", "8"),
    ("  4.5", "Convert", "8"),
    ("  4.6", "Sound Waves", "9"),
    ("  4.7", "Script Writer", "9"),
    ("  4.8", "Channel Identity", "9"),
    ("  4.9", "Scene Images", "9"),
    ("  4.10", "Footage", "10"),
    ("  4.11", "Mouse Automation", "10"),
    ("  4.12", "Automation Editor", "10"),
    ("  4.13", "API Storage", "10"),
    ("  4.14", "Brave Automation", "11"),
    ("  4.15", "Translate", "11"),
    ("  4.16", "Jobs Centre", "11"),
    ("  4.17", "Notifications", "11"),
    ("5", "Improvement Roadmap", "12"),
    ("6", "Known Limitations", "14"),
    ("7", "Troubleshooting", "14"),
]
tbl = doc.add_table(rows=len(toc_items), cols=3)
tbl.style = "Table Grid"
for row_idx, (num, title, page) in enumerate(toc_items):
    cells = tbl.rows[row_idx].cells
    cells[0].text = num
    cells[1].text = title
    cells[2].text = page
    for ci, cell in enumerate(cells):
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        if row_idx % 2 == 0:
            set_cell_bg(cell, "F5F3FF")
        else:
            set_cell_bg(cell, "FAFAFF")
# column widths
for row in tbl.rows:
    row.cells[0].width = Cm(1.5)
    row.cells[1].width = Cm(12)
    row.cells[2].width = Cm(1.5)

add_page_break(doc)


# ════════════════════════════════════════════════════════════════════════════════
# PAGE 3 — OVERVIEW
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, "1. Overview & Purpose", 1)
body(doc,
    "WeJaWi Studio is a cross-platform desktop application (Windows, macOS including "
    "Apple Silicon M-series) built with Python and PyQt5. It consolidates a suite of "
    "video-production automation tools into a single window, removing the need to "
    "switch between multiple standalone programs during content creation workflows."
)
body(doc,
    "The application is designed for YouTubers, podcasters, and short-form video creators "
    "who need fast, offline, GPU-accelerated processing without cloud subscriptions."
)

heading(doc, "2. System Requirements", 1)
tbl2 = doc.add_table(rows=8, cols=2)
tbl2.style = "Table Grid"
reqs = [
    ("Operating System", "macOS 12 + (Apple Silicon or Intel)  |  Windows 10/11 64-bit  |  Ubuntu 22+"),
    ("Python", "3.10 or higher"),
    ("FFmpeg", "v5 + installed and available in PATH (brew install ffmpeg on macOS)"),
    ("RAM", "8 GB minimum, 16 GB recommended for large files"),
    ("GPU (optional)", "Apple M-series ANE/Metal  |  NVIDIA CUDA  |  AMD ROCm/AMF"),
    ("Disk", "1 GB for application; additional space for model downloads and temp files"),
    ("Display", "1280 × 800 minimum; HiDPI/Retina supported"),
]
set_cell_bg(tbl2.rows[0].cells[0], "7C3AED")
set_cell_bg(tbl2.rows[0].cells[1], "7C3AED")
for ci, hdr in enumerate(("Component", "Requirement")):
    run = tbl2.rows[0].cells[ci].paragraphs[0].add_run(hdr)
    run.font.color.rgb = WHITE
    run.font.bold = True
for i, (k, v) in enumerate(reqs):
    tbl2.rows[i + 1].cells[0].text = k
    tbl2.rows[i + 1].cells[1].text = v
    if i % 2 == 0:
        for c in tbl2.rows[i + 1].cells:
            set_cell_bg(c, "F5F3FF")

add_page_break(doc)


# ════════════════════════════════════════════════════════════════════════════════
# PAGE 4 — FEATURE SET
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, "3. Feature Set", 1)
body(doc,
    "WeJaWi Studio provides 17 tools accessible from the left sidebar. "
    "Each tool opens in the main content area without leaving the application."
)

features = [
    ("🌐  Browse", "Home dashboard. Displays quick-launch cards for every tool. Click a card to jump directly to that tool."),
    ("🎙  Transcribe",
     "Converts speech in video/audio files to timestamped text segments.\n"
     "• Backend priority on Apple Silicon: MLX-Whisper (Metal + Neural Engine) → faster-whisper CPU → openai-whisper MPS → CPU.\n"
     "• Backend priority on Windows/Linux: faster-whisper CUDA → DirectML → openai-whisper CUDA → CPU.\n"
     "• Model sizes: tiny (fastest), base, small.\n"
     "• Outputs SRT, VTT, JSON, or plain text.\n"
     "• Supports 7 languages: English, German, Polish, Swedish, Norwegian, Hungarian, Spanish."),
    ("💬  Captions",
     "Burns subtitles directly into a video file using FFmpeg and the ASS subtitle format.\n"
     "• Live preview pane with adjustable font, size, text colour, and stroke colour.\n"
     "• Accepts .srt, .vtt, or plain .txt transcripts; also runs auto-transcription before render.\n"
     "• Encoder options: NVIDIA NVENC, AMD AMF, CPU libx264 (auto-detected).\n"
     "• Resizable split pane; right control panel is scrollable."),
    ("✂  Stitch Up",
     "Concatenates multiple video clips into a single output file.\n"
     "• Drag-and-drop clips from Finder/Explorer into the clip list.\n"
     "• Reorder by dragging rows or using the Move Up / Move Down buttons.\n"
     "• Mode: Re-encode (universal compatibility) or Fast Copy (no quality loss, requires matching codecs).\n"
     "• Shows file duration next to each clip name.\n"
     "• Progress bar and FFmpeg log during render."),
    ("🔁  Convert",
     "Batch file format conversion (video, audio, image). Configurable codec, resolution, and quality presets."),
    ("🎚  Sound Waves",
     "Audio waveform visualisation and processing. Applies audio filters, normalisation, or exports waveform images."),
    ("✍  Script Writer",
     "AI-assisted or template-based script drafting for YouTube videos, podcasts, or social posts."),
    ("💫  Channel Identity",
     "Manages branding assets — logo, colour palette, font choices — and applies them consistently across exports."),
    ("🎬  Scene Images",
     "Extracts key frames / scene thumbnails from a video at configurable intervals or on scene-change detection."),
    ("📼  Footage",
     "Organises raw footage files, provides quick preview, and lets you tag or annotate clips for later use."),
    ("🖱  Mouse Automation",
     "Records and replays mouse-and-keyboard macros. Useful for automating repetitive GUI workflows."),
    ("⚙  Automation Editor",
     "Visual editor for building multi-step automation sequences without writing code."),
    ("🔑  API Storage",
     "Securely stores API keys (YouTube, OpenAI, etc.) for use by other tools. Keys are saved locally in the app config directory."),
    ("🦁  Brave Automation",
     "Launches and controls Brave Browser for web-based automation tasks (channel uploads, scheduling, etc.)."),
    ("🌍  Translate",
     "Translates transcript text or captions into multiple languages, optionally using local or cloud models."),
    ("📋  Jobs Centre",
     "Queue view for all running and completed background tasks. Click a completed job to reveal its output file in Finder/Explorer."),
    ("🔔  Notifications",
     "In-app notification history. On macOS and Windows, completed jobs also trigger a native system tray notification."),
]

for name, desc in features:
    heading(doc, name, 2)
    body(doc, desc)

add_page_break(doc)


# ════════════════════════════════════════════════════════════════════════════════
# PAGE 5 — SOPs
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, "4. Standard Operating Procedures", 1)


def sop_section(title, steps, notes=None):
    heading(doc, title, 2)
    for i, s in enumerate(steps, 1):
        step(doc, i, s)
    if notes:
        p = doc.add_paragraph()
        run = p.add_run("Note: ")
        run.font.bold = True
        run.font.color.rgb = PURPLE
        p.add_run(notes)


sop_section("4.1  Browse (Home)", [
    "Launch WeJaWi Studio.",
    "The Browse page opens automatically, showing all tool cards.",
    "Click any card to navigate to that tool.",
    "Use the sidebar on the left for direct navigation at any time.",
])

sop_section("4.2  Transcribe", [
    "Click 🎙 Transcribe in the sidebar.",
    "Click Browse… and select a video or audio file (MP4, MOV, MKV, MP3, WAV, etc.).",
    "Choose Language from the dropdown.",
    "Select a Model: tiny for speed, base/small for accuracy.",
    "Select Device: on M-series Macs choose 'Metal / Neural Engine' for maximum speed.",
    "Click Transcribe. The status log shows progress; the first run downloads the selected model (~150 MB for tiny).",
    "When complete, the transcript appears in the text area.",
    "Click Save As… to export as .srt, .vtt, .json, or .txt.",
], notes="FFmpeg must be installed (brew install ffmpeg). Model files are cached after first download.")

sop_section("4.3  Captions", [
    "Click 💬 Captions in the sidebar.",
    "Under '1) Choose video', click Browse… and select your video file.",
    "Click Load / Snapshot. The first frame appears in the preview pane.",
    "Under '2) Style', choose font, caption size, text colour (white recommended), stroke colour (black recommended), and stroke size.",
    "Under '3) Transcript', either: (a) click 'Load transcript' and pick an .srt or .vtt file, (b) paste text in the text area, or (c) leave blank and enable 'Auto-transcribe'.",
    "If using Auto-transcribe, set Language, Model, and Device in the STT row.",
    "Under '4) Render', choose Video encoder (Auto recommended) and set the output file path.",
    "Click 'Render captions → video'. Monitor progress in the log.",
    "When done, a dialog shows the saved file path.",
])

sop_section("4.4  Stitch Up", [
    "Click ✂ Stitch Up in the sidebar.",
    "Add clips by clicking 'Add clips…' or dragging video files directly into the clip list from Finder.",
    "Reorder clips by dragging rows up/down, or select a clip and use the ▲/▼ buttons.",
    "To remove clips, select them and click Remove (or Clear all to start fresh).",
    "In the Output panel, choose File path for the merged video.",
    "Select Mode: 'Re-encode' works with any mixed sources; 'Fast copy' is faster but requires clips to share the same codec.",
    "Click 'Stitch & render'. Progress and FFmpeg output appear in the log at the bottom.",
    "A dialog confirms the saved path when complete.",
], notes="Re-encode mode with CPU libx264 is the safest option when clips come from different cameras or apps.")

sop_section("4.5  Convert", [
    "Click 🔁 Convert in the sidebar.",
    "Add source files.",
    "Choose the output format and quality preset.",
    "Click Convert and monitor the Jobs Centre for progress.",
])

sop_section("4.6  Sound Waves", [
    "Click 🎚 Sound Waves in the sidebar.",
    "Load an audio or video file.",
    "Select the processing preset (normalise, compress, EQ, waveform image).",
    "Click Process and save the output.",
])

sop_section("4.7  Script Writer", [
    "Click ✍ Script Writer in the sidebar.",
    "Enter topic/title and choose a template or tone.",
    "Click Generate (uses local logic or connected API key from API Storage).",
    "Edit the script in the text pane and export as .txt or .docx.",
])

sop_section("4.8  Channel Identity", [
    "Click 💫 Channel Identity in the sidebar.",
    "Upload your logo and specify brand colours and font.",
    "Apply the identity to export templates or overlays.",
])

sop_section("4.9  Scene Images", [
    "Click 🎬 Scene Images in the sidebar.",
    "Load a video file.",
    "Set the extraction interval or choose 'Scene change detection'.",
    "Click Extract. Thumbnails are saved to the chosen output folder.",
])

sop_section("4.10  Footage", [
    "Click 📼 Footage in the sidebar.",
    "Browse or drag-drop raw footage files.",
    "Click any clip to preview. Add tags or notes in the metadata panel.",
    "Use the filter bar to find clips by tag.",
])

sop_section("4.11  Mouse Automation", [
    "Click 🖱 Mouse Automation in the sidebar.",
    "Press Record and perform your mouse/keyboard actions.",
    "Press Stop when done.",
    "Optionally set a repeat count, then press Play to replay the macro.",
])

sop_section("4.12  Automation Editor", [
    "Click ⚙ Automation Editor in the sidebar.",
    "Drag action blocks from the palette onto the canvas.",
    "Configure each block's parameters in the property panel.",
    "Click Run to execute the sequence; view results in the log.",
])

sop_section("4.13  API Storage", [
    "Click 🔑 API Storage in the sidebar.",
    "Click Add Key, enter the service name and key value.",
    "Keys are stored in ~/Library/Application Support/WeJaWi/settings.json (macOS) or equivalent.",
    "Other tools read keys automatically from this store.",
])

sop_section("4.14  Brave Automation", [
    "Ensure Brave Browser is installed in /Applications.",
    "Click 🦁 Brave Automation in the sidebar.",
    "Enter the target URL or choose a saved automation preset.",
    "Click Run Automation. Brave opens and the sequence executes.",
])

sop_section("4.15  Translate", [
    "Click 🌍 Translate in the sidebar.",
    "Paste or load the text / subtitle file to translate.",
    "Select source and target languages.",
    "Click Translate. Review and export the result.",
])

sop_section("4.16  Jobs Centre", [
    "Click 📋 Jobs Centre in the sidebar at any time.",
    "Active jobs show a spinner and progress percentage.",
    "Completed jobs show a ✓ and the output file path.",
    "Click the path to reveal the file in Finder (macOS) or Explorer (Windows).",
])

sop_section("4.17  Notifications", [
    "Click 🔔 Notifications in the sidebar.",
    "View a timestamped list of all recent system events.",
    "Native tray notifications also appear when jobs complete (macOS, Windows).",
])

add_page_break(doc)


# ════════════════════════════════════════════════════════════════════════════════
# PAGE 6 — IMPROVEMENT ROADMAP
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, "5. Improvement Roadmap", 1)
body(doc,
    "The items below are prioritised enhancements identified during initial macOS porting. "
    "Priority levels: 🔴 High · 🟡 Medium · 🟢 Nice-to-have."
)

improvements = [
    ("🔴", "Browse page horizontal overflow",
     "Quick-action buttons and PINS column are cut off on widths below 1100 px. "
     "Fix by converting the card grid to a responsive FlowLayout or wrapping the area in a scroll view."),
    ("🔴", "Captions — macOS subtitle filter path escaping",
     "The ASS subtitle FFmpeg filter requires special colon escaping on macOS. "
     "Test and harden the _ff_filter_escape_for_subtitles() function with paths "
     "that contain spaces and parentheses, which are common in macOS home directories."),
    ("🔴", "FFmpeg not-found UX",
     "Several tools silently fail if FFmpeg is missing. Replace with a first-run "
     "check that shows a one-click install guide (brew install ffmpeg) on macOS."),
    ("🟡", "Stitch Up — clip thumbnail strip",
     "Show a small video thumbnail beside each clip name in the list for easier visual identification."),
    ("🟡", "Stitch Up — trim handles",
     "Add start/end trim fields per clip so users can cut head/tail before stitching without a separate tool."),
    ("🟡", "Transcribe — live progress",
     "MLX-Whisper and faster-whisper emit segment-level callbacks. Surface these as a real progress "
     "bar instead of a spinner, especially useful for long files."),
    ("🟡", "Transcribe — output format chooser",
     "Currently implicit. Add explicit SRT / VTT / TXT / JSON radio buttons so users know what they'll get."),
    ("🟡", "Captions — position control",
     "Allow captions to be placed at top, middle, or a custom vertical offset, not just bottom-centre."),
    ("🟡", "Captions — word-level highlight",
     "Optionally colour the currently spoken word a different colour (karaoke style) using ASS {\\k} tags."),
    ("🟡", "Convert — GPU encoder on macOS",
     "Add VideoToolbox (Apple's hardware H.264/H.265 encoder) as an encoder option. "
     "FFmpeg supports it via -c:v h264_videotoolbox."),
    ("🟡", "Script Writer — API key integration",
     "Wire the OpenAI key from API Storage into Script Writer so GPT-4o generation works without re-entering credentials."),
    ("🟡", "Jobs Centre — cancel individual jobs",
     "The cancel button currently only cancels the most recent job. Add per-row cancel support."),
    ("🟡", "Settings — output directory preference",
     "Add a global default output folder setting so users don't have to re-pick the path for every render."),
    ("🟢", "Dark / Light mode animation",
     "Smooth colour transition (150 ms) when toggling themes instead of an instant repaint."),
    ("🟢", "Sidebar — collapsible",
     "Allow the sidebar to collapse to icon-only width (~52 px) for users on smaller screens."),
    ("🟢", "Drag-and-drop into tool pages",
     "Accept file drops anywhere on a tool page (not just within specific file-picker areas)."),
    ("🟢", "Keyboard shortcuts",
     "Add Cmd/Ctrl+O (open file) and Cmd/Ctrl+R (render/run) to all tool pages for faster operation."),
    ("🟢", "Notification sound options",
     "Let users choose a custom notification sound or disable sound per-tool."),
    ("🟢", "Thumbnail / icon for the dock",
     "Supply a proper .icns (macOS) and .ico (Windows) icon so the app shows a branded icon in the Dock/Taskbar."),
    ("🟢", "Auto-update",
     "Background version check against a releases endpoint; prompt user when a new version is available."),
    ("🟢", "Bundled distribution",
     "Package with PyInstaller or Briefcase so users can run the app without a Python environment."),
]

tbl3 = doc.add_table(rows=len(improvements) + 1, cols=3)
tbl3.style = "Table Grid"
headers = ("Priority", "Feature / Area", "Description")
set_cell_bg(tbl3.rows[0].cells[0], "7C3AED")
set_cell_bg(tbl3.rows[0].cells[1], "7C3AED")
set_cell_bg(tbl3.rows[0].cells[2], "7C3AED")
for ci, h in enumerate(headers):
    run = tbl3.rows[0].cells[ci].paragraphs[0].add_run(h)
    run.font.color.rgb = WHITE
    run.font.bold = True
    run.font.size = Pt(10)

for ri, (pri, feat, desc) in enumerate(improvements):
    row = tbl3.rows[ri + 1]
    row.cells[0].text = pri
    row.cells[1].text = feat
    row.cells[2].text = desc
    bg = "FFF5F5" if "🔴" in pri else ("FFFBF0" if "🟡" in pri else "F0FFF4")
    for cell in row.cells:
        set_cell_bg(cell, bg)
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)

# column widths
for row in tbl3.rows:
    row.cells[0].width = Cm(1.8)
    row.cells[1].width = Cm(5.5)
    row.cells[2].width = Cm(9.5)

add_page_break(doc)


# ════════════════════════════════════════════════════════════════════════════════
# PAGE 7 — KNOWN LIMITATIONS & TROUBLESHOOTING
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, "6. Known Limitations", 1)
limits = [
    "FFmpeg must be installed separately — it is not bundled with the application.",
    "Transcription model files (MLX-Whisper, faster-whisper) are downloaded on first use; internet access is required for initial setup.",
    "Brave Automation requires Brave Browser installed in the default system location.",
    "Mouse Automation macros are not portable across different screen resolutions or OS scaling factors.",
    "The Convert and Sound Waves pages are functional stubs; some presets may not yet be fully implemented.",
    "Live mic recording is not supported — Transcribe and Captions operate on existing files only.",
    "Multi-language caption rendering may require additional fonts installed on the host OS.",
]
for lim in limits:
    bullet(doc, lim)

heading(doc, "7. Troubleshooting", 1)
issues = [
    ("App won't launch",
     "Ensure Python 3.10+ is installed and all dependencies are present. "
     "Run: pip3 install -r requirements.txt"),
    ("FFmpeg errors during render",
     "Verify FFmpeg is in PATH: run ffmpeg -version in Terminal. "
     "On macOS: brew install ffmpeg"),
    ("Transcription produces no output",
     "Check the log for [stt] lines. On first run, wait for the model download to complete. "
     "Ensure the input file has an audio track."),
    ("Captions page subtitle filter fails",
     "Avoid special characters (colons, parentheses) in the output path. "
     "Place input/output files in a simple directory like ~/Desktop."),
    ("Stitch output video has audio sync issues",
     "Switch Mode to 'Re-encode' instead of 'Fast copy'. "
     "Ensure all source clips have the same frame rate."),
    ("Brave not found",
     "Install Brave Browser to /Applications on macOS. "
     "On Windows, ensure it is in Program Files or Program Files (x86)."),
    ("Theme doesn't update correctly",
     "Toggle Settings → Dark mode / Light mode and restart the application."),
]
tbl4 = doc.add_table(rows=len(issues) + 1, cols=2)
tbl4.style = "Table Grid"
set_cell_bg(tbl4.rows[0].cells[0], "7C3AED")
set_cell_bg(tbl4.rows[0].cells[1], "7C3AED")
for ci, h in enumerate(("Problem", "Solution")):
    run = tbl4.rows[0].cells[ci].paragraphs[0].add_run(h)
    run.font.color.rgb = WHITE
    run.font.bold = True

for ri, (prob, sol) in enumerate(issues):
    tbl4.rows[ri + 1].cells[0].text = prob
    tbl4.rows[ri + 1].cells[1].text = sol
    if ri % 2 == 0:
        for c in tbl4.rows[ri + 1].cells:
            set_cell_bg(c, "F5F3FF")
    for c in tbl4.rows[ri + 1].cells:
        c.paragraphs[0].runs[0].font.size = Pt(9.5)

for row in tbl4.rows:
    row.cells[0].width = Cm(6)
    row.cells[1].width = Cm(10.5)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════════════════════
# BACK COVER
# ════════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("\n\n\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("WeJaWi Studio")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = PURPLE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Internal document — not for distribution")
run.font.size = Pt(10)
run.font.color.rgb = MID_GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Generated {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
run.font.size = Pt(9)
run.font.color.rgb = MID_GREY


# ── save ──────────────────────────────────────────────────────────────────────
doc.save(DOC_OUT)
print(f"Saved → {DOC_OUT}")
